# helper.py
from __future__ import annotations
import os
import requests
import time
from typing import Optional, Dict
from sqlalchemy.exc import OperationalError, DisconnectionError
from app.keyvault import get_secret
from docx import Document
from flask import send_file, Response
from datetime import datetime
import tempfile
from azure.identity import DefaultAzureCredential
from app import db


# -----------------------------
# Public configuration loaders
# -----------------------------
def load_config() -> dict[str, str | dict]:
    """
    Load core application secrets.
    Prefers environment variables; falls back to AKV if available.
    Returns a plain dict; nothing is cached globally.
    """
    # C7 / CH / NameAPI
    c7_key = get_secret("C7APIKey")
    c7_userid = get_secret("C7USERID")
    ch_key = get_secret("CHKEY")
    sql_username = get_secret("SQL_USERNAME", "SQL-USERNAME")
    sql_password = get_secret("SQL_PASSWORD", "SQL-PASSWORD")
    prefix = get_secret("NAMEAPI-KEYPREFIX")
    suffix = get_secret("NAMEAPI-KEYSUFFIX")
    nameapi_key = f"{prefix}-{suffix}"

    c7_hdr = {
        "Cache-Control": "no-cache",
        "Ocp-Apim-Subscription-Key": c7_key,
    }

    return {
        "C7_KEY": c7_key,
        "C7_USERID": c7_userid,
        "CH_KEY": ch_key,
        "NAMEAPI_KEY": nameapi_key,
        "C7_HDR": c7_hdr,
        "SQL_USERNAME": sql_username,
        "SQL_PASSWORD": sql_password,
    }


def load_azure_app_identity() -> Dict[str, str]:
    """
    Load Azure app-registration credentials (when not using managed identity).
    These should normally be provided by the platform as env vars.
    """
    return {
        "AZURE_CLIENT_ID": get_secret("AZURE-CLIENT-ID"),
        "AZURE_TENANT_ID": get_secret("AZURE-TENANT-ID"),
        "AZURE_CLIENT_SECRET": get_secret("AZURE-CLIENT-SECRET"),
    }


# -----------------------------
# Utilities
# -----------------------------
def formatName(name_string: str) -> str:
    """
    Convert 'Surname, Forename[:extra]' -> 'Forename Surname'
    """
    name_array = name_string.split(",")
    forename_part = name_array[1].strip() if len(name_array) > 1 else ""
    surname_part = name_array[0].strip()
    surname_only = surname_part.split(":")[0].strip()
    return f"{forename_part} {surname_only}".strip()


def uploadToSharePoint(file_bytes: bytes, filename: str, target_url):
    """
    Upload a file to SharePoint using Microsoft Graph API and managed identity.
    """    

    if debugMode():
        print(f"{datetime.now().strftime('%H:%M:%S')} uploadToSharePoint: Uploading file '{filename}' to SharePoint at '{target_url}'")

    # Configure region for Azure - explicitly set region for better performance
    credential = DefaultAzureCredential(
        additionally_allowed_tenants=["*"],
        # Add exclude options to speed up credential resolution
        exclude_visual_studio_code_credential=True,
        exclude_shared_token_cache_credential=True,
        exclude_powershell_credential=True
    )

    token = credential.get_token("https://graph.microsoft.com/.default")
    access_token = token.token

    site_name = get_secret('SP-SITE-NAME')
    site_domain = get_secret('SP-SITE-DOMAIN')
    library = get_secret('SP-LIBRARY')
    folder_path = target_url
    file_name = filename

    # Get Site ID
    site_url = f'https://graph.microsoft.com/v1.0/sites/{site_domain}:/sites/{site_name}'
    site_response = requests.get(site_url, headers={'Authorization': f'Bearer {access_token}'})

    if site_response.status_code != 200:
        print(f"Error getting site ID: {site_response.status_code} - {site_response.text}")
        return None

    site_id = site_response.json()['id']

    # Upload file
    upload_url = f'https://graph.microsoft.com/v1.0/sites/{site_id}/drive/root:/{library}/{folder_path}/{file_name}:/content'
    print(f"Uploading to URL: {upload_url}")

    upload_response = requests.put(upload_url, headers={
        'Authorization': f'Bearer {access_token}',
        'Content-Type': 'application/octet-stream'
    }, data=file_bytes)

    return upload_response.status_code


def downloadFromSharePoint(folder_path: str, filename: str) -> Optional[bytes]:
    """
    Download a file from SharePoint using Microsoft Graph API and managed identity.
    Returns the file bytes if successful, else None.
    """
    # Configure region for Azure - explicitly set region for better performance
    credential = DefaultAzureCredential(
        additionally_allowed_tenants=["*"],
        # Add exclude options to speed up credential resolution
        exclude_visual_studio_code_credential=True,
        exclude_shared_token_cache_credential=True,
        exclude_powershell_credential=True
    )

    token = credential.get_token("https://graph.microsoft.com/.default")
    access_token = token.token

    site_name = get_secret('SP-SITE-NAME')
    site_domain = get_secret('SP-SITE-DOMAIN')
    library= "Common"

    # Get Site ID
    site_url = f'https://graph.microsoft.com/v1.0/sites/{site_domain}:/sites/{site_name}'
    site_response = requests.get(site_url, headers={'Authorization': f'Bearer {access_token}'})

    if site_response.status_code != 200:
        print(f"Error getting site ID: {site_response.status_code} - {site_response.text}")
        return None

    site_id = site_response.json()['id']

    # Get file metadata first to get the download URL
    metadata_url = f'https://graph.microsoft.com/v1.0/sites/{site_id}/drive/root:/{library}/{folder_path}/{filename}'
    print(f"Getting file metadata from: {metadata_url}")

    metadata_response = requests.get(metadata_url, headers={
        'Authorization': f'Bearer {access_token}'
    })

    if metadata_response.status_code == 200:
        metadata = metadata_response.json()
        download_url = metadata.get('@microsoft.graph.downloadUrl')
        
        if download_url:
            print(f"Downloading file content from: {download_url}")
            
            # Download the actual file content using the download URL
            file_response = requests.get(download_url)
            
            if file_response.status_code == 200:
                content = file_response.content
                
                # Check if content looks like a valid DOCX file (should start with PK)
                if len(content) > 0 and content[:2] == b'PK':
                    if debugMode():
                        print(f"Downloaded file size: {len(content)} bytes")
                    return content
                else:
                    print(f"Downloaded content is not a valid DOCX file, size: {len(content)} bytes")
                    print(f"Content preview: {content[:100]}")
                    return None
            else:
                print(f"Error downloading file content: {file_response.status_code}")
                return None
        else:
            print("No download URL found in metadata")
            return None
    else:
        print(f"Error getting file metadata: {metadata_response.status_code} - {metadata_response.text}")
        return None


def wait_for_db(max_wait=120, interval=5):
    """Wait for the Azure database to be available before starting the app."""
    waited = 0
    while waited < max_wait:
        try:
            # Check if the connection is alive
            db.session.connection()
            print("Database is available.")
            return True
        except OperationalError:
            print(f"Database not available, retrying in {interval}s...")
            time.sleep(interval)
            waited += interval
    print("Database did not become available in time.")
    return False


def debugMode():
    config_mode = os.environ.get('FLASK_CONFIG', 'DevelopmentConfig')
    return config_mode == 'DevelopmentConfig'


def serve_docx(file_bytes: bytes, filename: str, replacements: Optional[dict] = None):
    """
    Open a docx from bytes, replace placeholders, convert to PDF and serve for viewing
    """
    import tempfile
    import os
    from flask import Response
    
    # Default replacements if none provided
    if replacements is None:
        replacements = {}
    
    # Add today's date as default replacement
    today_str = datetime.today().strftime('%d %B %Y')
    replacements.setdefault('{{DocDate}}', today_str)
    
    # Validate input
    if not file_bytes or len(file_bytes) < 100:
        raise ValueError(f"Invalid file_bytes provided, size: {len(file_bytes) if file_bytes else 0}")
    
    # Check if it looks like a DOCX file
    if file_bytes[:2] != b'PK':
        raise ValueError("File does not appear to be a valid DOCX file (missing PK header)")
    
    tmp_dir = tempfile.gettempdir()
    tmp_path = None
    modified_path = None
    pdf_path = None

    try:
        # Create unique temporary file paths
        import uuid
        unique_id = str(uuid.uuid4())
        tmp_path = os.path.join(tmp_dir, f"docx_input_{unique_id}.docx")
        modified_path = os.path.join(tmp_dir, f"docx_modified_{unique_id}.docx")
        pdf_path = os.path.join(tmp_dir, f"pdf_output_{unique_id}.pdf")

        # Write bytes to temp file
        with open(tmp_path, 'wb') as f:
            f.write(file_bytes)

        if debugMode():
            print(f"Created temp file: {tmp_path}")
            print(f"File exists: {os.path.exists(tmp_path)}")
            print(f"File size: {os.path.getsize(tmp_path)} bytes")

        # Open and modify the document
        doc = Document(tmp_path)
        if debugMode():
            print("Document loaded successfully")
            
        # Replace placeholders in the document
        replace_text_in_document(doc, replacements)

        # Save the modified document to a new file
        doc.save(modified_path)

        if debugMode():
            print(f"Document modifications completed and saved to: {modified_path}")
            print(f"Modified file exists: {os.path.exists(modified_path)}")
            print(f"Modified file size: {os.path.getsize(modified_path)} bytes")

        # Convert the MODIFIED DOCX to PDF
        convert_docx_to_pdf(modified_path, pdf_path)

        if debugMode():
            print("Document converted to PDF successfully")
            print(f"PDF file exists: {os.path.exists(pdf_path)}")

        # Read the PDF file and return it
        with open(pdf_path, 'rb') as pdf_file:
            pdf_content = pdf_file.read()

        # Return PDF response for inline viewing
        return Response(
            pdf_content,
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'inline; filename="{filename}.pdf"'
            }
        )

    except Exception as e:
        if debugMode():
            print(f"Error in serve_docx: {str(e)}")
            import traceback
            traceback.print_exc()
        raise  # Re-raise the exception so Flask can handle it properly
        
    finally:
        # Clean up temp files
        try:
            if debugMode():
                print(f"Cleaning up temp files")
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)
            if modified_path and os.path.exists(modified_path):
                os.unlink(modified_path)
            if pdf_path and os.path.exists(pdf_path):
                os.unlink(pdf_path)
        except Exception as cleanup_error:
            if debugMode():
                print(f"Cleanup error: {str(cleanup_error)}")


def convert_docx_to_pdf(docx_path: str, pdf_path: str):
    """
    Convert a DOCX file to PDF using available conversion methods
    """
    try:
        # Use python-docx with reportlab (basic conversion)
        try:
            convert_docx_to_pdf_reportlab(docx_path, pdf_path)
            if debugMode():
                print("Converted using reportlab")
            return
        except Exception as e:
            if debugMode():
                print(f"reportlab conversion failed: {e}")
        
        # If all methods fail, raise an error
        raise Exception("No PDF conversion method available. Please install docx2pdf, pypandoc, or ensure Microsoft Word is available.")
        
    except Exception as e:
        raise Exception(f"Failed to convert DOCX to PDF: {str(e)}")


def convert_docx_to_pdf_reportlab(docx_path: str, pdf_path: str):
    """
    Convert DOCX to PDF using reportlab (simple text-based conversion)
    """
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.colors import black
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        if debugMode():
            print(f"Converting DOCX to PDF: {docx_path} -> {pdf_path}")
        
        # Open the document
        doc = Document(docx_path)
        
        # Create a simple canvas-based PDF
        c = canvas.Canvas(pdf_path, pagesize=A4)
        width, height = A4
        
        # Starting position
        y_position = height - 50  # Start near top with margin
        line_height = 20
        left_margin = 50
        right_margin = 50
        
        # Set default font
        c.setFont("Helvetica", 12)
        c.setFillColor(black)
        
        if debugMode():
            print(f"Processing {len(doc.paragraphs)} paragraphs")
        
        def get_alignment_from_paragraph(paragraph):
            """Get text alignment from Word paragraph"""
            try:
                alignment = paragraph.alignment
                if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    return 'center'
                elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    return 'right'
                elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    return 'justify'
                else:
                    return 'left'  # Default and LEFT alignment
            except:
                return 'left'  # Fallback
        
        def draw_text_with_wrapping_and_alignment(text, start_y, alignment='left'):
            """Helper function to draw text with proper line wrapping, newline handling, and alignment"""
            current_y = start_y
            max_width = width - left_margin - right_margin
            
            # Split by explicit newlines first
            lines = text.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    # Empty line - just add vertical space
                    current_y -= line_height
                    if current_y < 50:
                        c.showPage()
                        c.setFont("Helvetica", 12)
                        c.setFillColor(black)
                        current_y = height - 50
                    continue
                
                # Handle word wrapping for each line
                words = line.split(' ')
                current_line = ""
                
                for word in words:
                    # Test if adding this word would exceed the line
                    test_line = current_line + (" " if current_line else "") + word
                    text_width = c.stringWidth(test_line, "Helvetica", 12)
                    
                    if text_width <= max_width:
                        current_line = test_line
                    else:
                        # Write the current line with proper alignment
                        if current_line:
                            draw_aligned_text(c, current_line, current_y, alignment, left_margin, right_margin, max_width)
                            current_y -= line_height
                            
                            # Check if we need a new page
                            if current_y < 50:
                                c.showPage()
                                c.setFont("Helvetica", 12)
                                c.setFillColor(black)
                                current_y = height - 50
                        
                        current_line = word
                
                # Write the last line of this paragraph line with proper alignment
                if current_line:
                    draw_aligned_text(c, current_line, current_y, alignment, left_margin, right_margin, max_width)
                    current_y -= line_height
                    
                    # Check if we need a new page
                    if current_y < 50:
                        c.showPage()
                        c.setFont("Helvetica", 12)
                        c.setFillColor(black)
                        current_y = height - 50
            
            return current_y
        
        def draw_aligned_text(canvas, text, y_pos, alignment, left_margin, right_margin, max_width):
            """Draw text with specified alignment"""
            text_width = canvas.stringWidth(text, "Helvetica", 12)
            
            if alignment == 'center':
                x_pos = left_margin + (max_width - text_width) / 2
            elif alignment == 'right':
                x_pos = width - right_margin - text_width
            elif alignment == 'justify':
                # For justify, we'll just use left alignment for now
                # True justification would require spacing words differently
                x_pos = left_margin
            else:  # left alignment (default)
                x_pos = left_margin
            
            canvas.drawString(x_pos, y_pos, text)
        
        # Process each paragraph as simple text with alignment
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            if text.strip():
                # Get paragraph alignment
                alignment = get_alignment_from_paragraph(paragraph)
                
                if debugMode() and i < 10:  # Only log first 10 paragraphs
                    text_preview = text[:100].replace('\n', '\\n')
                    print(f"Paragraph {i} ({alignment}): {text_preview}...")
                
                y_position = draw_text_with_wrapping_and_alignment(text, y_position, alignment)
                
                # Add extra space after paragraphs
                y_position -= 5
                
                # Check if we need a new page
                if y_position < 50:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    c.setFillColor(black)
                    y_position = height - 50
        
        # Process tables as simple text (tables will use left alignment)
        for table_idx, table in enumerate(doc.tables):
            if debugMode():
                print(f"Processing table {table_idx}")
            
            y_position -= 10  # Extra space before table
            
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    # Handle newlines in table cells, replace with space
                    cell_text = cell.text.strip().replace('\n', ' ')
                    row_texts.append(cell_text)
                
                # Use tab characters to separate table cells
                row_text = "\t".join(row_texts)
                if row_text.strip():
                    # Truncate long table rows and handle wrapping
                    if len(row_text) > 120:  # Increased limit for tab-separated content
                        row_text = row_text[:117] + "..."
                    
                    y_position = draw_text_with_wrapping_and_alignment(row_text, y_position, 'left')
                    
                    if y_position < 50:
                        c.showPage()
                        c.setFont("Helvetica", 12)
                        c.setFillColor(black)
                        y_position = height - 50
        
        # Save the PDF
        c.save()
        
        if debugMode():
            print(f"PDF created successfully: {os.path.exists(pdf_path)}")
            print(f"PDF file size: {os.path.getsize(pdf_path)} bytes")
        
    except ImportError:
        raise Exception("reportlab not available for PDF conversion")
    except Exception as e:
        if debugMode():
            print(f"Error in convert_docx_to_pdf_reportlab: {e}")
            import traceback
            traceback.print_exc()
        raise


def replace_text_in_document(doc, replacements: dict):
    """
    Simple text replacement throughout the document
    """
    if debugMode():
        print(f"Starting document replacement with {len(replacements)} replacements")
    
    # Simple paragraph-based replacement
    for i, paragraph in enumerate(doc.paragraphs):
        original_text = paragraph.text
        new_text = original_text
        
        # Apply all replacements
        for placeholder, replacement in replacements.items():
            new_text = new_text.replace(placeholder, replacement)
        
        # If text changed, replace the entire paragraph
        if new_text != original_text:
            # Clear all runs and add new text
            for run in paragraph.runs[:]:
                run._element.getparent().remove(run._element)
            
            # Add new run with replacement text, preserving newlines
            new_run = paragraph.add_run(new_text)
            
            if debugMode():
                # Fix: Move the string replacement outside the f-string
                original_preview = original_text[:50].replace('\n', '\\n')
                new_preview = new_text[:50].replace('\n', '\\n')
                print(f"Replaced paragraph {i}: '{original_preview}...' -> '{new_preview}...'")
    
    # Simple table replacement
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    new_text = original_text
                    
                    # Apply all replacements
                    for placeholder, replacement in replacements.items():
                        new_text = new_text.replace(placeholder, replacement)
                    
                    # If text changed, replace the paragraph
                    if new_text != original_text:
                        # Clear all runs and add new text
                        for run in paragraph.runs[:]:
                            run._element.getparent().remove(run._element)
                        
                        # Add new run with replacement text
                        paragraph.add_run(new_text)
                        
                        if debugMode():
                            # Fix: Move the string replacement outside the f-string
                            original_preview = original_text[:30].replace('\n', '\\n')
                            new_preview = new_text[:30].replace('\n', '\\n')
                            print(f"Replaced table cell: '{original_preview}...' -> '{new_preview}...'")


def execute_db_query_with_retry(stmt, operation_name="database query"):
    """
    Execute a database query with retry logic for connection issues.
    
    Args:
        stmt: SQLAlchemy statement to execute
        operation_name: Description of the operation for logging
    
    Returns:
        Query results or empty list if all retries fail
    """
    max_retries = 3
    retry_delay = 1  # seconds
    
    for attempt in range(max_retries):
        try:
            query_result = db.session.execute(stmt).scalars().all()
            return query_result
            
        except (OperationalError, DisconnectionError) as e:
            if debugMode():
                print(f"{datetime.now().strftime('%H:%M:%S')} {operation_name}: Database connection error on attempt {attempt + 1}: {str(e)}")
            
            if attempt < max_retries - 1:
                # Close the current session to ensure clean retry
                try:
                    db.session.rollback()
                except:
                    pass
                
                time.sleep(retry_delay)
                retry_delay *= 2  # Exponential backoff
                continue
            else:
                if debugMode():
                    print(f"{datetime.now().strftime('%H:%M:%S')} {operation_name}: All retry attempts failed, returning empty list")
                return []
        
        except Exception as e:
            if debugMode():
                print(f"{datetime.now().strftime('%H:%M:%S')} {operation_name}: Unexpected error: {str(e)}")
            return []